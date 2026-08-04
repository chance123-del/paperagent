from __future__ import annotations

"""Formal delivery exports for the generated LaTeX manuscript."""

from html import escape as xml_escape
import re
import shutil
import subprocess
from pathlib import Path


def _plain_latex(value: str) -> str:
    value = re.sub(r"\\(?:textbf|textit|emph|underline)\{([^{}]*)\}", r"\1", value)
    value = re.sub(r"\\href\{[^{}]*\}\{([^{}]*)\}", r"\1", value)
    value = value.replace(r"\&", "&").replace(r"\%", "%").replace(r"\_", "_")
    value = re.sub(r"\\[a-zA-Z]+(?:\[[^\]]*\])?", "", value)
    return value.replace("{", "").replace("}", "").strip()


def _bibliography_entries(tex_path: Path, project_dir: Path) -> list[tuple[str, str]]:
    """Read rendered BibTeX entries when available, with a simple local fallback."""
    bbl_path = tex_path.with_suffix(".bbl")
    if bbl_path.exists():
        content = bbl_path.read_text(encoding="utf-8", errors="replace")
        matches = re.finditer(
            r"\\bibitem(?:\[[^\]]*\])?\{([^}]+)\}\s*(.*?)(?=\\bibitem|\\end\{thebibliography\})",
            content,
            re.DOTALL,
        )
        entries = [(match.group(1).strip(), _plain_latex(match.group(2))) for match in matches]
        if entries:
            return entries

    bibliography = re.search(r"\\bibliography\{([^}]+)\}", tex_path.read_text(encoding="utf-8", errors="replace"))
    if not bibliography:
        return []
    bib_path = project_dir / f"{bibliography.group(1).strip()}.bib"
    if not bib_path.exists():
        return []
    content = bib_path.read_text(encoding="utf-8", errors="replace")
    entries: list[tuple[str, str]] = []
    for match in re.finditer(r"(?im)^\s*@\w+\s*\{\s*([^,\s]+)\s*,(.*?)(?=^\s*@\w+\s*\{|\Z)", content, re.DOTALL):
        key, fields = match.group(1).strip(), match.group(2)
        values = {
            name.lower(): _plain_latex(value)
            for name, value in re.findall(r'(?im)^\s*(author|title|journal|publisher|year)\s*=\s*[{\"]([^}\"]+)[}\"]', fields)
        }
        parts = [values[name] for name in ("author", "title", "journal", "publisher", "year") if values.get(name)]
        if parts:
            entries.append((key, ". ".join(parts) + "."))
    return entries


def _word_text(value: str, citation_numbers: dict[str, int]) -> str:
    def replace_citation(match: re.Match[str]) -> str:
        keys = [key.strip() for key in match.group(1).split(",") if key.strip()]
        labels = [str(citation_numbers[key]) for key in keys if key in citation_numbers]
        return f"[{', '.join(labels)}]" if labels else ""

    value = re.sub(r"\\cite(?:[a-zA-Z*]+)?(?:\[[^\]]*\])?\{([^}]+)\}", replace_citation, value)
    return _plain_latex(value)


def _fallback_docx(tex_path: Path, output_path: Path, project_dir: Path) -> str:
    from docx import Document
    from docx.shared import Cm, Pt

    source = tex_path.read_text(encoding="utf-8", errors="replace")
    bibliography_entries = _bibliography_entries(tex_path, project_dir)
    citation_numbers = {key: number for number, (key, _) in enumerate(bibliography_entries, start=1)}
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    title = re.search(r"\\title\{(.+?)\}", source, re.DOTALL)
    if title:
        document.add_heading(_plain_latex(title.group(1)), level=0)

    in_abstract = False
    in_figure = False
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            text = _word_text(" ".join(paragraph), citation_numbers)
            if text:
                document.add_paragraph(text)
            paragraph.clear()

    for raw in source.splitlines():
        line = raw.strip()
        if not line or line.startswith("%"):
            flush()
            continue
        if line.startswith(r"\begin{abstract}"):
            flush()
            in_abstract = True
            document.add_heading("Abstract", level=1)
            continue
        if line.startswith(r"\end{abstract}"):
            flush()
            in_abstract = False
            continue
        heading = re.match(r"\\(section|subsection|subsubsection)\{(.+)\}", line)
        if heading:
            flush()
            document.add_heading(_word_text(heading.group(2), citation_numbers), level={"section": 1, "subsection": 2, "subsubsection": 3}[heading.group(1)])
            continue
        if line.startswith(r"\begin{figure}"):
            flush()
            in_figure = True
            continue
        image = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}", line)
        if image:
            candidate = project_dir / image.group(1)
            if candidate.exists():
                try:
                    document.add_picture(str(candidate), width=Cm(14))
                except Exception:
                    pass
            continue
        caption = re.search(r"\\caption\{(.+)\}", line)
        if caption:
            flush()
            document.add_paragraph(_plain_latex(caption.group(1)), style="Caption")
            continue
        if line.startswith(r"\end{figure}"):
            in_figure = False
            continue
        if line.startswith("\\") or in_figure:
            continue
        paragraph.append(line)

    flush()
    if bibliography_entries:
        document.add_heading("References", level=1)
        for number, (_, entry) in enumerate(bibliography_entries, start=1):
            document.add_paragraph(f"[{number}] {entry}")
    document.core_properties.title = _plain_latex(title.group(1)) if title else tex_path.stem
    document.save(output_path)
    return "Word exported with the built-in structured DOCX converter (Pandoc was not found)."


def export_docx_from_tex(tex_path: str | Path, output_path: str | Path, project_dir: str | Path) -> tuple[bool, str]:
    """Export a DOCX, preferring Pandoc while guaranteeing a local fallback."""
    tex = Path(tex_path).resolve()
    destination = Path(output_path).resolve()
    project = Path(project_dir).resolve()
    destination.parent.mkdir(parents=True, exist_ok=True)
    pandoc = shutil.which("pandoc")
    if pandoc:
        command = [pandoc, str(tex), "--from=latex", "--to=docx", f"--resource-path={project}", "-o", str(destination)]
        completed = subprocess.run(command, cwd=str(project), capture_output=True, text=True, encoding="utf-8", errors="replace")
        if completed.returncode == 0 and destination.exists() and destination.stat().st_size:
            return True, "Word exported with Pandoc."
        pandoc_note = (completed.stderr or completed.stdout).strip()
    else:
        pandoc_note = "Pandoc not found"
    try:
        return True, _fallback_docx(tex, destination, project)
    except Exception as exc:
        return False, f"Word export failed: {pandoc_note}; fallback error: {exc}"


def _pdf_inline(value: str, citation_numbers: dict[str, int]) -> str:
    """Convert a small, safe LaTeX inline subset to ReportLab paragraph markup."""
    link_tokens: dict[str, str] = {}

    def preserve_link(match: re.Match[str]) -> str:
        token = f"@@PDFLINK{len(link_tokens)}@@"
        url = xml_escape(match.group(1).strip(), quote=True)
        text = xml_escape(_plain_latex(match.group(2)), quote=False)
        link_tokens[token] = f'<link href="{url}" color="#175CD3"><u>{text}</u></link>'
        return token

    value = re.sub(r"\\href\{([^{}]+)\}\{([^{}]+)\}", preserve_link, value)
    value = _word_text(value, citation_numbers)
    value = xml_escape(value, quote=False)
    value = value.replace(r"\%", "%").replace(r"\&", "&amp;").replace(r"\_", "_")
    for token, markup in link_tokens.items():
        value = value.replace(token, markup)
    return value


def _fallback_pdf(tex_path: Path, output_path: Path, project_dir: Path) -> str:
    """Render a dependable, fully local PDF when a LaTeX distribution is unavailable."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.pdfmetrics import registerFont
    from reportlab.platypus import (
        Image,
        KeepTogether,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    registerFont(UnicodeCIDFont("STSong-Light"))
    source = tex_path.read_text(encoding="utf-8", errors="replace")
    bibliography_entries = _bibliography_entries(tex_path, project_dir)
    citation_numbers = {key: number for number, (key, _) in enumerate(bibliography_entries, start=1)}
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "PaperBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=17,
        textColor=colors.HexColor("#344054"),
        spaceAfter=7,
    )
    title_style = ParagraphStyle(
        "PaperTitle",
        parent=body,
        fontSize=20,
        leading=27,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#102A43"),
        spaceAfter=14,
    )
    heading_styles = {
        "section": ParagraphStyle("PaperH1", parent=body, fontSize=15, leading=21, textColor=colors.HexColor("#0B2E59"), spaceBefore=14, spaceAfter=7),
        "subsection": ParagraphStyle("PaperH2", parent=body, fontSize=12.5, leading=18, textColor=colors.HexColor("#155EEF"), spaceBefore=10, spaceAfter=5),
        "subsubsection": ParagraphStyle("PaperH3", parent=body, fontSize=11, leading=16, textColor=colors.HexColor("#344054"), spaceBefore=8, spaceAfter=4),
    }
    caption_style = ParagraphStyle("PaperCaption", parent=body, fontSize=9, leading=13, alignment=TA_CENTER, textColor=colors.HexColor("#667085"))
    formula_style = ParagraphStyle("PaperFormula", parent=body, fontName="Courier", fontSize=9, leading=13, alignment=TA_CENTER, backColor=colors.HexColor("#F2F4F7"), borderPadding=7)
    reference_style = ParagraphStyle("PaperReference", parent=body, fontSize=9, leading=14, leftIndent=12, firstLineIndent=-12)

    story = []
    title_match = re.search(r"\\title\{(.+?)\}", source, re.DOTALL)
    story.append(Paragraph(_pdf_inline(title_match.group(1), citation_numbers) if title_match else xml_escape(tex_path.stem), title_style))

    paragraph_lines: list[str] = []
    equation_lines: list[str] = []
    table_rows: list[list[str]] = []
    in_document = False
    in_equation = False
    in_tabular = False
    in_figure = False
    figure_parts: list[object] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            text = _pdf_inline(" ".join(paragraph_lines), citation_numbers)
            if text.strip():
                story.append(Paragraph(text, body))
            paragraph_lines.clear()

    def flush_table() -> None:
        if not table_rows:
            return
        width = max(len(row) for row in table_rows)
        normalized = [row + [""] * (width - len(row)) for row in table_rows]
        data = [[Paragraph(_pdf_inline(cell, citation_numbers), body) for cell in row] for row in normalized]
        table = Table(data, repeatRows=1, hAlign="CENTER")
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF2F8")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2E59")),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D5DD")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.extend([table, Spacer(1, 8)])
        table_rows.clear()

    for raw in source.splitlines():
        line = raw.strip()
        if line == r"\begin{document}":
            in_document = True
            continue
        if not in_document:
            continue
        if line == r"\end{document}":
            break
        if not line or line.startswith("%"):
            flush_paragraph()
            continue
        if line.startswith(r"\begin{equation") or line.startswith(r"\["):
            flush_paragraph()
            in_equation = True
            equation_lines = []
            continue
        if in_equation:
            if line.startswith(r"\end{equation") or line == r"\]":
                formula = xml_escape(" ".join(equation_lines), quote=False)
                story.extend([Paragraph(formula or " ", formula_style), Spacer(1, 7)])
                equation_lines = []
                in_equation = False
            else:
                equation_lines.append(line)
            continue
        if line.startswith(r"\begin{tabular}"):
            flush_paragraph()
            in_tabular = True
            table_rows = []
            continue
        if in_tabular:
            if line.startswith(r"\end{tabular}"):
                in_tabular = False
                flush_table()
            elif not line.startswith((r"\toprule", r"\midrule", r"\bottomrule", r"\hline")):
                cleaned = re.sub(r"\\\\\s*$", "", line)
                table_rows.append([_plain_latex(cell.strip()) for cell in cleaned.split("&")])
            continue
        heading = re.match(r"\\(section|subsection|subsubsection)\{(.+)\}", line)
        if heading:
            flush_paragraph()
            story.append(Paragraph(_pdf_inline(heading.group(2), citation_numbers), heading_styles[heading.group(1)]))
            continue
        if line.startswith(r"\begin{abstract}"):
            flush_paragraph()
            story.append(Paragraph("摘要 / Abstract", heading_styles["section"]))
            continue
        if line.startswith(r"\end{abstract}"):
            flush_paragraph()
            continue
        if line.startswith(r"\begin{figure}"):
            flush_paragraph()
            in_figure = True
            figure_parts = []
            continue
        image_match = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}", line)
        if image_match:
            relative = Path(image_match.group(1))
            candidates = [tex_path.parent / relative, project_dir / relative, project_dir / relative.name]
            candidate = next((item for item in candidates if item.exists()), None)
            if candidate:
                try:
                    rendered = Image(str(candidate))
                    rendered._restrictSize(16 * cm, 10 * cm)
                    figure_parts.append(rendered)
                except Exception:
                    figure_parts.append(Paragraph(f"图像资源：{xml_escape(relative.name)}", caption_style))
            continue
        caption_match = re.search(r"\\caption\{(.+)\}", line)
        if caption_match:
            figure_parts.append(Spacer(1, 4))
            figure_parts.append(Paragraph(_pdf_inline(caption_match.group(1), citation_numbers), caption_style))
            continue
        if line.startswith(r"\end{figure}"):
            in_figure = False
            if figure_parts:
                story.extend([KeepTogether(figure_parts), Spacer(1, 9)])
            figure_parts = []
            continue
        if in_figure:
            continue
        if line.startswith(
            (
                r"\maketitle",
                r"\author",
                r"\date",
                r"\bibliography",
                r"\bibliographystyle",
                r"\begin{table}",
                r"\end{table}",
                r"\centering",
                r"\label",
            )
        ):
            continue
        paragraph_lines.append(line)

    flush_paragraph()
    if bibliography_entries:
        story.append(Paragraph("参考文献 / References", heading_styles["section"]))
        for number, (_, entry) in enumerate(bibliography_entries, start=1):
            story.append(Paragraph(f"[{number}] {xml_escape(entry, quote=False)}", reference_style))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=2.2 * cm,
        leftMargin=2.2 * cm,
        topMargin=2.1 * cm,
        bottomMargin=2.0 * cm,
        title=_plain_latex(title_match.group(1)) if title_match else tex_path.stem,
        author="PaperFormat Agent",
    )

    def decorate_page(canvas, doc) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#D0D5DD"))
        canvas.setLineWidth(0.4)
        canvas.line(2.2 * cm, 1.45 * cm, A4[0] - 2.2 * cm, 1.45 * cm)
        canvas.setFont("STSong-Light", 8)
        canvas.setFillColor(colors.HexColor("#667085"))
        canvas.drawString(2.2 * cm, 1.0 * cm, "PaperFormat Agent · 本地兼容 PDF")
        canvas.drawRightString(A4[0] - 2.2 * cm, 1.0 * cm, f"第 {doc.page} 页")
        canvas.restoreState()

    document.build(story, onFirstPage=decorate_page, onLaterPages=decorate_page)
    return "PDF exported with the built-in local renderer; no network or LaTeX package download was required."


def export_pdf_from_tex(tex_path: str | Path, output_path: str | Path, project_dir: str | Path) -> tuple[bool, str]:
    """Export a readable PDF locally when native LaTeX compilation is unavailable."""
    tex = Path(tex_path).resolve()
    destination = Path(output_path).resolve()
    project = Path(project_dir).resolve()
    try:
        note = _fallback_pdf(tex, destination, project)
        return bool(destination.exists() and destination.stat().st_size), note
    except Exception as exc:
        return False, f"Built-in PDF export failed: {exc}"
