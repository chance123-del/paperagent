from __future__ import annotations

"""Convert authoring formats into a self-contained LaTeX project."""

from dataclasses import dataclass, field
from pathlib import Path
import re
import shutil
import unicodedata
import zipfile


@dataclass
class SourceDocument:
    title: str
    blocks: list[tuple[str, object]] = field(default_factory=list)
    abstract_text: str = ""
    keywords_text: str = ""
    images: list[Path] = field(default_factory=list)
    figure_captions: list[str] = field(default_factory=list)
    table_captions: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)


def _escape_latex(value: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$",
        "#": r"\#", "_": r"\_", "{": r"\{", "}": r"\}", "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in value)


def _clean_text(value: str) -> str:
    value = value.replace("\u00a0", " ").replace("\x00", " ")
    value = "".join(" " if unicodedata.category(char).startswith("C") else char for char in value)
    value = re.sub(r"\(cid:\d+\)", "[formula preserved below]", value)
    return re.sub(r"\s+", " ", value).strip()


def _split_long_text(value: str, maximum: int = 1600) -> list[str]:
    if len(value) <= maximum:
        return [value] if value else []
    chunks: list[str] = []
    remaining = value
    while len(remaining) > maximum:
        boundary = max(remaining.rfind(mark, 0, maximum) for mark in (". ", "? ", "! ", "; ", " "))
        boundary = boundary if boundary > maximum // 2 else maximum
        chunks.append(remaining[:boundary + 1].strip())
        remaining = remaining[boundary + 1:].strip()
    if remaining:
        chunks.append(remaining)
    return chunks


def _pdf_page_blocks(raw_text: str) -> list[tuple[str, object]]:
    lines = [_clean_text(line) for line in raw_text.replace("\r", "\n").split("\n")]
    lines = [line for line in lines if line]
    blocks: list[tuple[str, object]] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        text = _clean_text(" ".join(paragraph_lines))
        for chunk in _split_long_text(text):
            blocks.append(("paragraph", chunk))
        paragraph_lines = []

    known_headings = {
        "abstract", "introduction", "background", "methods", "methodology", "materials and methods",
        "results", "discussion", "conclusion", "conclusions", "references", "acknowledgments",
        "appendix", "\u6458\u8981", "\u5f15\u8a00", "\u65b9\u6cd5", "\u7ed3\u679c", "\u8ba8\u8bba", "\u7ed3\u8bba", "\u53c2\u8003\u6587\u732e",
    }
    boilerplate = (
        "available online", "published by", "all rights reserved", "creativecommons.org", "license", "copyright", "doi.org",
        "corresponding author", "e-mail addresses", "contents lists available", "journal homepage", "received ", "accepted ", "sciencedirect",
        "contributed equally", "journal of radiation research and applied sciences",
    )
    for line in lines:
        lowered = line.lower()
        if any(marker in lowered for marker in boilerplate) or re.fullmatch(r"\d+", line):
            continue
        inline_heading = re.match(r"^(\d+(?:\.\d+)*)\s*\.\s*(introduction|background|methods?|methodology|materials? and methods?|results?(?: and discussion)?|discussion|conclusions?|references|acknowledg(?:e)?ments?|appendix)\s+(.+)$", line, re.IGNORECASE)
        if inline_heading:
            if paragraph_lines:
                flush_paragraph()
            level = min(inline_heading.group(1).count(".") + 1, 3)
            blocks.append(("heading", (level, inline_heading.group(2).title())))
            paragraph_lines.append(inline_heading.group(3))
            continue
        numbered = re.match(r"^(\d+(?:\.\d+)*)\.?\s+(.+)$", line)
        heading_text = numbered.group(2).strip().lower() if numbered else lowered
        known_prefixes = ("introduction", "background", "method", "material", "result", "discussion", "conclusion", "reference", "acknowledg", "appendix")
        is_heading = heading_text in known_headings or (numbered is not None and heading_text.startswith(known_prefixes))
        if is_heading and paragraph_lines:
            flush_paragraph()
        if is_heading:
            level = min(numbered.group(1).count(".") + 1, 3) if numbered else 1
            blocks.append(("heading", (level, numbered.group(2).strip() if numbered else line)))
        else:
            paragraph_lines.append(line)
    flush_paragraph()
    return blocks


def _contains_cjk(document: SourceDocument) -> bool:
    values = [document.title, document.abstract_text, document.keywords_text]
    for _, payload in document.blocks:
        values.append(str(payload))
    return any("\u4e00" <= char <= "\u9fff" for value in values for char in value)


def _words_to_lines(words: list[dict]) -> str:
    ordered = sorted(words, key=lambda word: (round(float(word["top"]) / 3), float(word["x0"])))
    lines: list[list[dict]] = []
    for word in ordered:
        if not lines or abs(float(word["top"]) - float(lines[-1][0]["top"])) > 3:
            lines.append([word])
        else:
            lines[-1].append(word)
    return "\n".join(" ".join(item["text"] for item in sorted(line, key=lambda word: float(word["x0"]))) for line in lines)


def _column_ordered_page_text(page) -> str:
    words = page.extract_words(x_tolerance=1, y_tolerance=3)
    usable = [word for word in words if 58 < float(word["top"]) < float(page.height) - 28]
    midpoint = float(page.width) / 2
    left = [word for word in usable if float(word["x0"]) < midpoint - 8]
    right = [word for word in usable if float(word["x0"]) >= midpoint - 8]
    if len(left) < 20 or len(right) < 20:
        return _words_to_lines(usable)
    return _words_to_lines(left) + "\n" + _words_to_lines(right)


def _extract_formula_crops(layout_pages, assets_dir: Path) -> dict[int, list[Path]]:
    formula_assets: dict[int, list[Path]] = {}
    for page_number, page in enumerate(layout_pages, start=1):
        bad_chars = [
            char for char in page.chars
            if str(char.get("text", "")).startswith("(cid:") and str(char.get("tag", "")).lower() == "formula"
        ]
        used_tops: list[float] = []
        for index, char in enumerate(bad_chars, start=1):
            top = float(char["top"])
            if any(abs(top - previous) < 14 for previous in used_tops):
                continue
            used_tops.append(top)
            midpoint = float(page.width) / 2
            left, right = (35, midpoint - 8) if float(char["x0"]) < midpoint else (midpoint + 8, float(page.width) - 35)
            crop = page.crop((left, max(45, top - 14), right, min(float(page.height) - 28, float(char["bottom"]) + 16)))
            destination = assets_dir / f"page-{page_number}-formula-{index}.png"
            crop.to_image(resolution=180).save(destination, format="PNG")
            formula_assets.setdefault(page_number, []).append(destination)
    return formula_assets


def _caption_kind_and_text(text: str) -> tuple[str, str] | None:
    match = re.match(r"^(figure|fig\.?|图|table|表)\s*\d*\s*[:：.\-]?\s*(.+)$", text, re.IGNORECASE)
    if not match:
        return None
    kind = "figure" if match.group(1).lower() in {"figure", "fig.", "fig", "图"} else "table"
    return kind, match.group(2).strip()


def _caption_text(value: str, fallback: str) -> str:
    cleaned = re.sub(r"^(?:figure|fig\.?|table|\u56fe|\u8868)\s*\d*\s*[:\uff1a.\-]?\s*", "", value, flags=re.IGNORECASE).strip()
    return cleaned or fallback


def _copy_docx_images(source: Path, assets_dir: Path) -> list[Path]:
    assets_dir.mkdir(parents=True, exist_ok=True)
    images: list[Path] = []
    with zipfile.ZipFile(source) as archive:
        for member in archive.namelist():
            if not member.startswith("word/media/") or member.endswith("/"):
                continue
            destination = assets_dir / Path(member).name
            with archive.open(member) as image_in, destination.open("wb") as image_out:
                shutil.copyfileobj(image_in, image_out)
            images.append(destination)
    return images


def load_docx(source: Path, assets_dir: Path) -> SourceDocument:
    from docx import Document

    document = Document(source)
    images = _copy_docx_images(source, assets_dir)
    blocks: list[tuple[str, object]] = []
    figure_captions: list[str] = []
    table_captions: list[str] = []
    title = source.stem
    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        caption = _caption_kind_and_text(text)
        if style == "caption" or caption:
            if caption and caption[0] == "table":
                document_caption = caption[1]
                # Store source captions separately so they become table captions, not body prose.
                # The sequential association is deterministic for a converted manuscript.
                table_captions.append(document_caption)
            else:
                figure_captions.append(caption[1] if caption else text)
            continue
        if style == "title" and title == source.stem:
            title = text
        elif "heading" in style:
            match = re.search(r"(\d+)", style)
            blocks.append(("heading", (min(int(match.group(1)) if match else 1, 3), text)))
        else:
            blocks.append(("paragraph", text))

    for table in document.tables:
        rows = [[_clean_text(cell.text) for cell in row.cells] for row in table.rows]
        if rows:
            blocks.append(("table", rows))
    return SourceDocument(title=title, blocks=blocks, images=images, figure_captions=figure_captions, table_captions=table_captions)


def load_markdown(source: Path, assets_dir: Path) -> SourceDocument:
    text = source.read_text(encoding="utf-8", errors="replace")
    title = source.stem
    blocks: list[tuple[str, object]] = []
    image_pattern = re.compile(r"!\[[^\]]*\]\(([^)]+)\)")
    images: list[Path] = []
    raw_lines = text.splitlines()
    index = 0
    while index < len(raw_lines):
        raw = raw_lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        image_match = image_pattern.fullmatch(line)
        if image_match:
            original = (source.parent / image_match.group(1)).resolve()
            if original.exists() and original.is_file():
                assets_dir.mkdir(parents=True, exist_ok=True)
                copied = assets_dir / original.name
                shutil.copy2(original, copied)
                images.append(copied)
            index += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level, value = len(heading.group(1)), heading.group(2).strip()
            if level == 1 and not blocks:
                title = value
            else:
                blocks.append(("heading", (level, value)))
            index += 1
            continue
        if "|" in line and index + 1 < len(raw_lines) and re.fullmatch(r"[|:\-\s]+", raw_lines[index + 1].strip()):
            rows = [[cell.strip() for cell in line.strip("|").split("|")]]
            index += 2  # Skip the Markdown header separator row.
            while index < len(raw_lines) and "|" in raw_lines[index]:
                row = raw_lines[index].strip()
                if not row:
                    break
                rows.append([cell.strip() for cell in row.strip("|").split("|")])
                index += 1
            blocks.append(("table", rows))
            continue
        elif not re.match(r"^[-*+]\s+", line):
            blocks.append(("paragraph", line))
        index += 1
    return SourceDocument(title=title, blocks=blocks, images=images)


def load_pdf(source: Path, assets_dir: Path) -> SourceDocument:
    from pypdf import PdfReader
    import pdfplumber

    reader = PdfReader(str(source))
    page_texts: list[str] = []
    images: list[Path] = []
    page_images: dict[int, list[Path]] = {}
    assets_dir.mkdir(parents=True, exist_ok=True)
    with pdfplumber.open(source) as layout_pdf:
        layout_texts = [_column_ordered_page_text(page) for page in layout_pdf.pages]
        formula_assets = _extract_formula_crops(layout_pdf.pages, assets_dir)
    for page_number, page in enumerate(reader.pages, start=1):
        raw_content = page.extract_text() or ""
        page_texts.append(raw_content if page_number == 1 else layout_texts[page_number - 1])
        for image_number, image in enumerate(page.images, start=1):
            suffix = Path(image.name).suffix or ".png"
            destination = assets_dir / f"page-{page_number}-image-{image_number}{suffix}"
            try:
                destination.write_bytes(image.data)
                images.append(destination)
                page_images.setdefault(page_number, []).append(destination)
            except OSError:
                pass
    for page_number, assets in formula_assets.items():
        page_images.setdefault(page_number, []).extend(assets)
    full_text = "\n".join(page_texts)
    introduction = re.search(r"(?im)^\s*1\s*\.\s*introduction\s*$", full_text)
    front_matter = full_text[:introduction.start()] if introduction else page_texts[0] if page_texts else ""
    title_lines = [_clean_text(line) for line in front_matter.splitlines() if _clean_text(line)]
    title_parts = []
    for line in title_lines[:3]:
        if re.match(r"^[A-Z][a-z]+\s+[A-Z][a-z]+", line) and title_parts:
            break
        if line.upper() in {"ARTICLE INFO", "ABSTRACT", "KEYWORDS:"}:
            break
        title_parts.append(line)
    title = _clean_text(" ".join(title_parts)) or source.stem
    abstract_match = re.search(r"(?is)\babstract\s*(.+?)(?=\n\s*1\s*\.\s*introduction\b)", full_text)
    keywords_match = re.search(r"(?is)\bkeywords?\s*:\s*(.+?)(?=\n\s*abstract\b)", front_matter)
    abstract_text = _clean_text(abstract_match.group(1)) if abstract_match else ""
    keywords_text = _clean_text(keywords_match.group(1)) if keywords_match else ""
    blocks: list[tuple[str, object]] = []
    for page_number, page_text in enumerate(page_texts, start=1):
        if page_number == 1 and introduction:
            page_introduction = re.search(r"(?im)^\s*1\s*\.\s*introduction\s*$", page_text)
            page_text = page_text[page_introduction.start():] if page_introduction else ""
        blocks.extend(_pdf_page_blocks(page_text))
        if page_number > 1:
            blocks.extend(("figure", image) for image in page_images.get(page_number, []))
    notes = []
    if not blocks:
        notes.append("PDF contains no extractable text. It may be a scan and requires OCR before reliable conversion.")
    return SourceDocument(title=title, blocks=blocks, abstract_text=abstract_text, keywords_text=keywords_text, images=images, notes=notes)


def load_source(source: Path, assets_dir: Path) -> SourceDocument:
    suffix = source.suffix.lower()
    if suffix == ".docx":
        return load_docx(source, assets_dir)
    if suffix in {".md", ".markdown"}:
        return load_markdown(source, assets_dir)
    if suffix == ".pdf":
        return load_pdf(source, assets_dir)
    raise ValueError("Supported manuscript formats are .docx, .pdf, .md, .tex and .zip.")


def render_latex(document: SourceDocument, rules: dict) -> str:
    geometry = rules.get("geometry", "a4paper,margin=2.54cm")
    line_spread = rules.get("line_spread", "1.5")
    documentclass = rules.get("documentclass", "article")
    documentclass_options = rules.get("documentclass_options", "12pt")
    lines = [
        r"\documentclass[" + documentclass_options + r"]{" + documentclass + r"}",
        r"\usepackage{graphicx}", r"\usepackage{booktabs}", r"\usepackage{float}",
        r"\usepackage{setspace}", r"\usepackage{hyperref}", r"\linespread{" + str(line_spread) + r"}\selectfont",
        r"\begin{document}", r"\title{" + _escape_latex(document.title) + r"}", r"\author{}", r"\date{}", r"\maketitle",
    ]
    if _contains_cjk(document):
        lines.insert(1, r"\usepackage[UTF8]{ctex}")
    if document.abstract_text:
        lines.extend([r"\begin{abstract}", _escape_latex(document.abstract_text), r"\end{abstract}"])
    if document.keywords_text:
        lines.extend([r"\noindent\textbf{Keywords: }" + _escape_latex(document.keywords_text), ""])
    if not rules.get("class_managed_layout"):
        lines.insert(1, r"\usepackage[" + geometry + r"]{geometry}")
    table_number = 0
    embedded_figures = False
    for kind, payload in document.blocks:
        if kind == "heading":
            level, text = payload
            command = {1: "section", 2: "subsection", 3: "subsubsection"}[level]
            lines.append(rf"\{command}{{{_escape_latex(text)}}}")
        elif kind == "paragraph":
            lines.extend([_escape_latex(str(payload)), ""])
        elif kind == "table":
            rows = payload
            if not rows:
                continue
            table_number += 1
            width = max(len(row) for row in rows)
            caption = document.table_captions[table_number - 1] if table_number <= len(document.table_captions) else "Imported table"
            lines.extend([r"\begin{table}[H]", r"\centering", rf"\caption{{{_escape_latex(_caption_text(caption, 'Imported table'))}}}", r"\begin{tabular}{" + "l" * width + "}", r"\toprule"])
            for row_index, row in enumerate(rows):
                values = [_escape_latex(value) for value in row] + [""] * (width - len(row))
                lines.append(" & ".join(values) + r" \\")
                if row_index == 0:
                    lines.append(r"\midrule")
            lines.extend([r"\bottomrule", r"\end{tabular}", r"\end{table}"])
        elif kind == "figure":
            image = payload
            embedded_figures = True
            caption = "Preserved formula" if "-formula-" in image.name else "Imported figure"
            lines.extend([
                r"\begin{figure}[H]", r"\centering", rf"\includegraphics[width=0.85\linewidth]{{assets/{image.name}}}",
                rf"\caption{{{caption}}}", r"\end{figure}",
            ])
    for image_number, image in enumerate([] if embedded_figures else document.images, start=1):
        caption = document.figure_captions[image_number - 1] if image_number <= len(document.figure_captions) else "Imported figure"
        lines.extend([
            r"\begin{figure}[H]", r"\centering", rf"\includegraphics[width=0.85\linewidth]{{assets/{image.name}}}",
            rf"\caption{{{_escape_latex(_caption_text(caption, 'Imported figure'))}}}", r"\end{figure}",
        ])
    lines.append(r"\end{document}")
    return "\n".join(lines) + "\n"
