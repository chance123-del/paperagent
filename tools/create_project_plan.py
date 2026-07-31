from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("项目书_论文排版与期刊交付智能体.docx")
BLUE = "1F4D78"
LIGHT = "E8EEF5"
CAUTION = "FFF2CC"
RED = "FCE4D6"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    tc_pr.append(node)


def borders(table, color="B7C9D6"):
    tbl_pr = table._tbl.tblPr
    border = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:color"), color)
        border.append(tag)
    tbl_pr.append(border)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9.5)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    borders(t)
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, bold=True, color="FFFFFF")
        shade(t.rows[0].cells[i], BLUE)
        if widths:
            t.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return t


def add_run(p, text, bold=False, color=None, size=10.5):
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def para(doc, text="", style=None, indent=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.28
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    add_run(p, text)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, size=10.2)
    return p


def note(doc, title, text, fill=CAUTION):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t, "D6B656" if fill == CAUTION else "D99A6C")
    cell = t.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    add_run(p, title + " ", bold=True, color="7A5A00" if fill == CAUTION else "9B1C1C")
    add_run(p, text)
    doc.add_paragraph()


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(5)
    add_run(p, text, bold=True, color=BLUE if level < 3 else "333333", size={1: 15, 2: 12.5, 3: 11}[level])
    return p


def main():
    d = Document()
    s = d.sections[0]
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(2.54); s.right_margin = Cm(2.54)
    normal = d.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(62)
    add_run(p, "论文排版与期刊交付智能体", bold=True, color=BLUE, size=24)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "项目计划书（可行性、范围、实施与验收）", color="404040", size=13)
    d.add_paragraph()
    table(d, ["项目属性", "约定"], [
        ("目标", "将作者已完成的论文初稿和附件，转换为可追溯、可核验的期刊/学校格式交付包。"),
        ("核心原则", "格式自动化可以做；正文事实、数据、结论、图表注、文献条目和链接不得由系统猜写。"),
        ("主要交付", "Word、PDF、LaTeX 源码/压缩包、格式核验报告、资产匹配清单与失败日志。"),
        ("适用边界", "优先支持文字型 DOCX/Markdown/LaTeX/PDF；扫描件必须先经 OCR 并标明不确定文本。"),
    ], [3.0, 13.2])
    para(d, "版本：V1.0（需求梳理稿）  |  日期：2026-07-31", style=None)
    d.add_page_break()

    heading(d, "一、项目定位与成功定义")
    para(d, "本项目不是论文代写工具，而是“作者内容保真 + 目标规则可追溯 + 多格式交付”的本地论文编排系统。用户上传已写好的文字版初稿、图表/附件压缩包、参考文献合集与格式规则后，系统完成解析、映射、排版、编译和核验；它不应改变作者的研究主张。", indent=True)
    table(d, ["成功标准", "可被验收的结果"], [
        ("内容保真", "正文段落、数值、引用键、图表文件和文献条目均可回溯到用户上传内容；无来源内容不进入最终稿。"),
        ("格式合规", "每一项已应用规则均能指出来源：官方模板/指南、用户自定义规则或内置通用规则。"),
        ("资产完整", "每个 [FigN]/[TableN] 的匹配状态明确为已匹配、缺失、重复、类型不符或待人工确认；替换后的版式遵循已确认的目标模板/期刊规则。"),
        ("交付可用", "LaTeX 工程可复编译；PDF 经过编译和检查；Word 为可编辑导出并说明转换限制。"),
    ], [3.2, 13.0])
    note(d, "关键提醒：", "“符合期刊要求”只能对已提供或已确认的规则负责。期刊名称的网络匹配只能帮助定位候选模板，不能把猜测当作官方规范。")

    heading(d, "二、用户、输入与输出")
    heading(d, "2.1 目标用户", 2)
    bullet(d, "作者：上传初稿、资料包和参考文献，选择期刊或上传学校/期刊规则。")
    bullet(d, "编辑/课题组助理：处理批量图表、检查缺项、复核待确认项并下载交付包。")
    bullet(d, "系统管理员：维护已核验期刊规则包、模板版本和允许的文件类型。")
    heading(d, "2.2 输入契约（上传即校验）", 2)
    table(d, ["输入", "允许形式", "系统要求与处理"], [
        ("论文初稿", "DOCX、Markdown、LaTeX；文字型 PDF", "保留原文件；抽取结构和正文。扫描 PDF 提示 OCR，OCR 不确定片段必须标记，不可静默替换。"),
        ("图表资料包", "ZIP（含图片、表格、可选清单）", "安全解压；拒绝路径穿越、可执行文件、超限文件；按文件名与占位符做精确匹配，并采集像素、DPI、纵横比和表格尺寸等元数据。"),
        ("图表注合集包", "单独 ZIP（内含 Excel/CSV/JSON 清单）", "不参与图表文件匹配；只按资产 ID 取回图注、表注、来源注、替代文本与可选链接，再由目标规则决定其位置。"),
        ("参考文献合集", "BibTeX 优先；RIS/EndNote/CSV 可导入转换", "保存原文件；生成解析报告；字段不足只报错/待补，不生成作者、题名、DOI。"),
        ("规则来源", "官方 LaTeX 模板、Author Guidelines、学校模板、JSON 规则、用户自定义", "记录来源、版本、上传时间、适用范围和证据页/段；冲突必须显示优先级。"),
        ("可选元数据", "项目配置表/表单", "图表题注、图注、超链接、替代文本、文件类型、作者确认记录。"),
    ], [2.4, 4.0, 9.8])
    heading(d, "2.3 输出交付包", 2)
    table(d, ["文件", "用途", "发布条件"], [
        ("main.tex + assets/", "主 LaTeX 与已引用资产", "无未决匹配和无阻断规则错误；保留固定标签。"),
        ("latex_source.zip", "可复编译工程", "含规则快照、references.bib、构建说明、资产清单；不含无关上传文件。"),
        ("main.pdf", "投稿/审阅版", "真实模板编译成功；若仅生成 article 预览，必须命名 preview.pdf，不能冒充正式稿。"),
        ("manuscript.docx", "可编辑交换稿", "导出完成并声明 Word 与 LaTeX 在浮动对象、交叉引用上的差异。"),
        ("format_report.md/json", "核验与审计", "列出规则状态、修复记录、警告、未解决事项、编译日志摘要。"),
        ("asset_manifest.csv/json", "图表/表格溯源", "列出占位符、原文件、哈希、题注来源、链接来源、输出标签、尺寸/分辨率、应用的版式规则和状态。"),
    ], [3.0, 5.2, 8.0])

    heading(d, "三、端到端工作流与交互")
    table(d, ["步骤", "系统动作", "用户决策/界面提醒"], [
        ("1. 新建项目", "建立隔离工作目录，计算上传文件哈希，生成 project.json。", "提示项目名称、目标期刊/学校、匿名稿要求及输出语言。"),
        ("2. 上传与解析", "抽取正文、标题层级、已有图表、图表占位符、引用占位符与文献；扫描 ZIP。", "显示解析覆盖率；提示扫描件 OCR、加密文件和未识别格式。"),
        ("3. 选择规则", "载入规则包并展示来源、版本、置信度与覆盖项。", "低置信度/未验证期刊规则不可一键发布，必须上传官方依据或改用通用规则。"),
        ("4. 匹配与预览", "对每个占位符建立候选资产；生成图表预览与替换差异。", "重复/缺失/类型不符必须由用户选择、补传或明确跳过。"),
        ("5. 编排", "只做格式变换、资产插入、图/表注安放、编号与引用渲染；生成中间 LaTeX。", "所有将改变正文文本的动作默认关闭；用户可逐项确认规则性修复。"),
        ("6. 编译与核验", "编译 PDF，执行规则、引用、资产、链接、版面和可复编译性检查。", "阻断项阻止正式交付；警告项以清单签收。"),
        ("7. 下载", "打包正式交付或明确标识的预览交付。", "展示版本号、输入哈希、规则快照和未关闭风险。"),
    ], [2.1, 7.0, 7.1])

    heading(d, "四、图表、图表注与超链接设计")
    heading(d, "4.1 推荐资料包规范", 2)
    para(d, "初稿中使用 [Fig1]、[Fig2]、[Table1] 等占位符。ZIP 根目录或子目录中使用同名资产，例如 Fig1.png、Fig2.tiff、Table1.xlsx。文件名规范只处理大小写、空格、下划线和 Figure/Fig 的等价归一化，不做模糊“猜配”。", indent=True)
    table(d, ["推荐文件", "作用", "是否必需"], [
        ("Fig1.png / Fig1.tiff", "图片素材；支持 PNG/JPG/TIFF，保留原始文件。", "图占位符存在时必需"),
        ("Table1.xlsx / Table1.csv", "表格数据源；导入为 LaTeX 表格或按规则转图片。", "表占位符存在时必需"),
        ("annotations.zip", "独立图表注合集包；内含 annotations.xlsx/CSV/JSON。", "推荐；图表规则要求题注/表注时必需"),
        ("annotations.xlsx", "Figures、Tables、Links 三个工作表；以资产 ID 关联。", "合集包内推荐格式"),
        ("manifest.csv", "资产 ID、文件、类型、版本、校验和、确认状态。", "批量/多人协作推荐"),
    ], [3.5, 9.1, 3.6])
    heading(d, "4.2 图表注合集包：与占位符严格分离", 2)
    para(d, "用户可单独上传 annotations.zip。该包是“文字和链接元数据”，不是图片/表格资源包，也不在正文中产生新的图表占位符。正文只保留 [Fig1]、[Table1] 等对象占位符；系统先匹配对象，再以 Fig1/Table1 作为键读取注释记录并附着到对应对象。", indent=True)
    table(d, ["工作表/文件", "最小字段", "系统行为"], [
        ("Figures", "asset_id, caption, note, source, alt_text", "仅接受 asset_id=FigN；匹配图像后写入题注/图注区。"),
        ("Tables", "asset_id, caption, note, source, alt_text", "仅接受 asset_id=TableN；匹配表格后写入表题/表注区。"),
        ("Links（可选）", "asset_id, url_or_doi, link_text", "与对应图表注绑定；用户只需粘贴 DOI 或 URL。"),
        ("JSON/CSV 等价格式", "字段名与上述一致", "便于自动化；界面同时提供模板下载和表格编辑。"),
    ], [3.4, 6.3, 6.5])
    note(d, "用户体验约定：", "不要求用户学习 LaTeX 或填写 \u005ccaption。用户只需在下载的 Excel 模板每行填写 Fig1/Table1 及其文字；未填写的图注/表注不会被系统凭空生成。")
    heading(d, "4.3 匹配算法与状态机", 2)
    table(d, ["规则", "行为", "结果"], [
        ("精确唯一匹配", "[Fig1] 对应归一化后唯一的 Fig1.png。", "自动插入，生成 \u005clabel{fig:1} 或规则指定标签。"),
        ("重复候选", "Fig1.png 与 Fig1.tiff 同时存在，或多份同名文件。", "停止自动选择，显示预览、分辨率/哈希和路径，要求用户确认。"),
        ("缺失资产", "初稿有 [Table3]，资料包无 Table3。", "保留原占位符并作为阻断项；不生成空表或替代内容。"),
        ("类型不符", "[Fig2] 只找到 .xlsx。", "标记错误，要求用户更正占位符或文件。"),
        ("无占位符资产", "ZIP 多出 Fig9.png。", "不自动插入；报告为未引用资产，供用户决定。"),
    ], [3.0, 8.0, 5.2])
    heading(d, "4.4 期刊图表版式规则与无损缩放", 2)
    para(d, "图表被匹配到占位符后，仍须通过“资产质量检查 + 目标模板规则渲染”两道关。版式不由当前段落的空白位置决定，而由目标期刊/学校模板规定的浮动位置、栏宽、最大尺寸、题注位置、编号和跨栏策略决定。", indent=True)
    table(d, ["版式项目", "规则来源与实现", "阻断/提醒"], [
        ("位置", "读取官方模板/规则中的 figure/table 浮动参数、允许位置、跨栏和图题/表题位置；使用模板宏实现。", "若规则未提供，使用通用默认并明确标为“非期刊已验证”。"),
        ("尺寸", "以单栏宽、双栏宽、最大高度、页边距为上限；保留原始像素、DPI 和纵横比。", "禁止横纵方向分别缩放；超过版心时提示用户提供更合适素材或允许等比缩小。"),
        ("分辨率/DPI", "按期刊对线稿、灰度图、彩色图的要求校验有效 DPI；不通过插值伪造高分辨率。", "分辨率不足为阻断或高优先级警告，取决于已确认期刊规则。"),
        ("表格布局", "根据模板字号、栏宽、最大列数和表题位置排版；优先换行、合理列宽或跨页 longtable。", "不通过压扁字体、裁剪单元格或拉伸表格来塞入版心。"),
        ("图像帧率", "静态图片和表格没有帧率；本项目按“分辨率/DPI”执行。补充视频应单独作为素材，记录 codec、帧率、时长和投稿平台要求。", "视频不直接替代论文中的静态图；无期刊规定时不自动转码。"),
    ], [2.5, 9.0, 4.7])
    bullet(d, "默认缩放策略：只允许等比缩放；渲染前后记录宽高、缩放比例和有效 DPI。任何裁剪、旋转、重采样或颜色空间转换都必须由用户显式选择并写入报告。")
    bullet(d, "输出 PDF 后进行视觉与规则核验：图表是否越界、是否跨栏正确、题注是否贴近对象、表格是否溢出、字体/线宽是否可读；检查失败不得只靠缩小到不可读来通过。")
    heading(d, "4.5 题注、图注、表注与交叉引用", 2)
    bullet(d, "题注（caption）来源优先级固定为：用户在独立 annotations.zip/表单明确填写 > 初稿中与占位符绑定的已存在题注 > 空值并报待补。系统不得根据图片内容自动写题注。")
    bullet(d, "图注/表注（note、Source、缩略语说明）独立保存，按期刊规则置于表下、图下或脚注区；它们不能与 caption 混为一项。")
    bullet(d, "期刊规则在图题/表题位置上优先于用户习惯：例如规则要求表题在表上方、图题在图下方时，系统自动安放对应记录；规则缺失时显示默认位置和待确认提示。")
    bullet(d, "首次成功插入后生成稳定标签，例如 fig:1、tab:1；正文中的“见图 1”应转为 \u005cref 或 \u005cautoref，禁止硬编码数字。")
    bullet(d, "编号由 LaTeX/官方模板生成；重新排序时自动更新。若规则要求图题在下、表题在上，规则引擎只调整位置，不改写文字。")
    heading(d, "4.6 超链接策略（低门槛输入）", 2)
    table(d, ["链接类型", "安全实现", "禁止行为"], [
        ("DOI/文献 URL", "只从 BibTeX、annotations.xlsx 的 Links 表或作者表单读取；使用 hyperref 生成可点击链接。", "根据题名搜索并猜填 DOI/URL。"),
        ("图表来源/数据仓库", "用户在 Links 表粘贴 DOI/URL，可选填显示文字；系统按 asset_id 绑定并自动生成规范链接。", "将图片像素识别结果、文件名或网页猜测为链接。"),
        ("文内跳转", "通过 \u005clabel + \u005cref/\u005chyperref 建立 PDF 内部锚点。", "用固定页码或手写编号模拟跳转。"),
        ("外部 URL", "校验协议、可访问性与显示文字；将失效检查记为警告。", "自动打开、上传或发布用户的私有链接。"),
    ], [3.0, 7.5, 5.7])
    bullet(d, "最简操作：在匹配结果行点击“添加链接”，粘贴完整 URL 或 DOI（如 10.xxxx/xxxx），系统规范化为 https://doi.org/...；不需要用户编写 LaTeX、HTML 或 Markdown。批量时只需粘贴到 Links 工作表。")
    note(d, "关键提醒：", "超链接既是内容也是合规证据。系统只负责嵌入和校验，不负责发现或臆造链接；链接是否可公开访问需要用户最终确认。")

    heading(d, "五、规则与参考文献治理")
    heading(d, "5.1 规则优先级与证据链", 2)
    table(d, ["优先级", "规则来源", "允许覆盖的范围", "可信级别"], [
        ("1", "目标期刊官方 LaTeX 模板 + 官方 Author Guidelines", "文档类、宏包、版面、图表、引文与投稿要求。", "已验证"),
        ("2", "用户上传的学校/导师/编辑部书面规则", "用户明确指定的局部格式。", "用户确认"),
        ("3", "已版本化的内置期刊规则包", "仅对有来源、测试和发布日期的规则项生效。", "已验证或待更新"),
        ("4", "通用基础规则", "无专用依据时的最小可编译、通用排版。", "通用，不得标为期刊合规"),
    ], [1.3, 5.4, 6.0, 3.5])
    para(d, "规则对象应至少记录 rule_id、字段、取值、来源文件、页码/章节、适用期刊/版本、优先级、置信状态和测试用例。系统的“期刊匹配”可使用公开元数据识别出版社或候选档案，但只在用户确认后绑定规则包。", indent=True)
    heading(d, "5.2 参考文献处理", 2)
    bullet(d, "以用户上传的 BibTeX 为权威源，原文件只读存档；内部可解析为统一条目模型，再导出 references.bib。")
    bullet(d, "仅可进行可逆、可记录的格式标准化，例如字段转义、引用样式切换、重复键报告；不补全缺失作者、页码、卷期、题名、DOI。")
    bullet(d, "正文引用只在已有引用键与文献条目之间连接。未解析引用、未被引用条目、重复 key、悬空 cite 都写入报告。")
    bullet(d, "期刊模板要求 BibTeX/Biber 时，构建链显式选择并连续编译至交叉引用稳定；构建日志必须保留。")
    heading(d, "5.3 文中引用占位符规则", 2)
    para(d, "与图表一致，正文引用应使用显式、可机器检查的占位符，推荐 [[cite:Smith2024]]、[[cite:Smith2024;Wang2023]] 或 [[citep:Smith2024]]。其中 key 必须与 BibTeX 条目键完全一致；cite/citep 等类型只映射到已确认的期刊引用命令。初稿中已有的 \\cite{...} 也应解析为同一内部引用对象。", indent=True)
    table(d, ["情形", "处理", "发布规则"], [
        ("唯一有效 key", "[[cite:Smith2024]] 转为模板要求的 \\cite{Smith2024}、\\citep{Smith2024} 或等价命令。", "允许自动转换并在 source mapping 留痕。"),
        ("多文献引用", "[[cite:Smith2024;Wang2023]] 按输入顺序转换为同一引用命令的多个 key。", "每个 key 必须存在；任何一个缺失即阻断。"),
        ("缺失或拼写不符", "显示候选 key 供用户选择；不擅自纠错，不删除占位符。", "阻断正式交付。"),
        ("无占位符的裸文献描述", "保留作者原文，不推断应引用哪一条文献。", "报告提示，不自动插入引用。"),
        ("未被引用文献", "保留在原始合集和审计报告中；是否进入最终参考文献表按期刊规则与用户选择。", "不自动删除。"),
    ], [3.0, 8.4, 4.8])
    note(d, "关键提醒：", "引用占位符只解决“把已有文献准确放入正文”的问题，不能证明该文献支持该论断。论断与文献的学术相关性仍由作者负责。")

    heading(d, "六、反编造与质量门禁（项目红线）")
    table(d, ["控制点", "实现要求", "发布影响"], [
        ("内容锁定", "保存输入哈希、段落/表格/资产来源 ID；默认仅转换格式，不调用改写或补全。", "正文差异没有对应规则或用户确认，阻断发布。"),
        ("证据闭环", "每一个题注、图注、链接、文献字段均记录来源文件和位置。", "来源为空时标为待补，不进入正式稿。"),
        ("最小自动化", "允许确定性转义、标签、编号、模板宏替换；禁止推断实验结果、解释图片、生成摘要/关键词。", "发现生成性内容标记，阻断正式交付。"),
        ("人工确认队列", "重复匹配、低置信规则、OCR 不确定文本、缺失题注、无效链接均进入队列。", "阻断项必须清零；警告项须导出时签收。"),
        ("可复现构建", "固定规则快照、模板版本、编译器版本和构建命令；生成 manifest。", "无法复编译的 PDF 只能作为失败/预览产物。"),
    ], [2.7, 8.0, 5.5])
    note(d, "发布门禁：", "正式交付必须同时满足：没有缺失/重复/类型不符的资产匹配；没有悬空引用；所有必需规则有依据；目标模板真实编译成功；核验报告中无阻断项。")

    heading(d, "七、基于现有原型的实施范围")
    para(d, "当前代码已具备初稿解析、规则处理、占位符扫描/插入、参考文献、LaTeX 编译、Word 导出和报告的基础模块。V1 应优先将这些能力从“可运行”提升为“可审计、可确认、可发布”，而不是先扩展大量期刊。", indent=True)
    table(d, ["现有基础", "V1 需补齐", "验收重点"], [
        ("[FigN]/[TableN] 扫描与 ZIP 文件名匹配", "加入 manifest、精确匹配状态、重复候选选择 UI、文件哈希、安全解压和未引用资产报告。", "重复时绝不默认选择；缺失时不产生空内容。"),
        ("图/表 LaTeX 插入", "加入独立 annotations.zip、notes、alt_text、稳定 label、交叉引用、Links 表和模板化超链接。", "题注、表注和链接均能追溯来源，且不与对象占位符混淆。"),
        ("内置期刊 profile 与指南解析", "规则证据模型、规则版本、冲突展示、模板优先与期刊确认页。", "不能把元数据匹配称作官方规则。"),
        ("BibTeX 与引用样式", "RIS/EndNote 导入、字段完整性检测、重复/悬空引用检查、BibTeX/Biber 构建策略。", "不补全任何书目信息。"),
        ("PDF/Word/源码导出", "正式/预览区分、交付清单、输出一致性核验、失败日志和可复现构建元数据。", "预览 PDF 绝不伪装为正式模板 PDF。"),
    ], [3.4, 8.3, 4.5])

    heading(d, "八、建议架构与数据对象")
    table(d, ["层", "职责", "关键对象"], [
        ("接入层", "上传、病毒/大小/类型检查、项目隔离、原件存档。", "Project, Upload, FileHash"),
        ("解析层", "提取论文结构、引用、占位符、文献和资料包索引。", "ManuscriptIR, Placeholder, Asset, BibliographyEntry"),
        ("规则层", "装载模板、规则优先级、证据定位、冲突与版本控制。", "RuleSet, RuleEvidence, RuleDecision"),
        ("确认层", "展示匹配候选、题注/链接来源、阻断项和用户签收。", "ReviewTask, Resolution, Approval"),
        ("编排层", "构造 LaTeX AST/项目，插入图表、引用与受控超链接。", "RenderPlan, SourceMapping"),
        ("核验与交付层", "编译、解析日志、检查 PDF/源文件、生成报告和 zip。", "ValidationResult, DeliveryManifest"),
    ], [2.5, 7.0, 6.7])
    para(d, "建议将 project.json 作为审计根文件：记录输入哈希、规则快照、用户选择、生成器版本、资产映射、构建命令和输出哈希。所有报告都从同一份清单派生，避免“界面显示”和“最终包”不一致。", indent=True)

    heading(d, "九、分期计划与里程碑")
    table(d, ["阶段", "范围", "完成标志"], [
        ("P0：需求冻结（1-2 周）", "确定首批支持的输入格式、一个或两个目标期刊/学校规则、资料包规范和阻断项。", "示例项目、规则来源、验收用例和非目标清单由业务方签字。"),
        ("P1：可信导入（2-3 周）", "项目存档、输入哈希、ZIP 安全扫描、占位符/资产索引、BibTeX 解析和基础报告。", "可准确报告每个占位符及每项参考文献的状态。"),
        ("P2：图表、引用与确认工作台（3-4 周）", "匹配预览、重复选择、annotations.zip、图表版式预检、题注/图注/表注/交叉引用及 [[cite:key]] 引用占位符。", "所有图表、图表注、链接和文中引用均可人工确认并生成审计清单。"),
        ("P3：规则与渲染（3-4 周）", "规则证据模型、官方模板接入、LaTeX 项目构建、PDF/Word/zip 输出。", "首批模板真实编译，正式/预览交付明确分流。"),
        ("P4：核验与试运行（2-3 周）", "回归测试、异常用例、版面检查、用户验收、日志和可用性修订。", "通过约定样本集；阻断门禁和报告可复核。"),
    ], [3.4, 8.2, 4.6])
    note(d, "范围控制提醒：", "首版应锁定“文字型初稿 + 规范 ZIP + BibTeX + 少量已核验模板”。OCR、复杂 Word 浮动对象、任意期刊网页自动抓取和自动补全书目应列为后续能力。")

    heading(d, "十、验收测试集与指标")
    table(d, ["测试类别", "样例", "通过标准"], [
        ("正常路径", "DOCX + [Fig1]/[Table1] + 唯一资产 + 完整 BibTeX + 已验证模板。", "生成四类交付物；资产、编号、引用、规则和 PDF 均通过。"),
        ("匹配异常", "缺失 Fig、重复 Fig、图片/表格类型错、未引用资产。", "无自动臆选；全部出现在确认队列和报告；正式发布被阻止。"),
        ("内容保真", "含数值、公式、引用、已有题注和中英文混排的初稿。", "输出映射可定位原始来源；无新增事实性文本。"),
        ("规则冲突", "用户规则与官方模板不同，或期刊名仅元数据匹配。", "显示冲突和证据；未经确认不按低可信规则发布。"),
        ("构建异常", "缺类文件、缺宏包、BibTeX 错误、超时。", "保留失败日志；不把降级预览标记为正式 PDF。"),
        ("安全与鲁棒", "恶意 ZIP、超大文件、路径穿越、损坏图片、非法 URL。", "隔离/拒绝并报告原因，原项目与宿主环境不受影响。"),
    ], [2.6, 8.0, 5.6])
    para(d, "建议指标：占位符唯一匹配准确率 100%（在规范命名样本中）；阻断项漏报率 0；正式交付的可复编译率 100%；每个输出图表、题注、链接和文献条目均有来源记录；首批人工确认完成时间和失败原因作为体验指标持续观测。", indent=True)

    heading(d, "十一、项目风险与处置")
    table(d, ["风险", "后果", "处置"], [
        ("期刊规则不完整或更新", "错误格式提交。", "版本化规则与证据；超过有效期提示复核；只对已验证项声明合规。"),
        ("Word/PDF 结构差异", "浮动图表、脚注、公式转换损失。", "明确源格式支持矩阵；保留原件；关键对象以 LaTeX 交付为准并做样本回归。"),
        ("图表命名混乱", "错图或漏图。", "要求 manifest；无唯一匹配不自动插入；预览确认。"),
        ("生成模型幻觉", "虚假数据、题注、文献或链接。", "内容锁定、来源字段必填、禁止无证据生成、发布门禁。"),
        ("LaTeX 依赖缺失", "PDF 无法生成。", "容器化/固定编译器；记录版本；失败只产出源码和日志。"),
        ("版权和隐私", "资料泄露或不当传播。", "默认本地处理、项目隔离、删除策略、外链不自动访问/上传。"),
    ], [3.0, 4.5, 8.7])

    heading(d, "十二、上线前必须确认的决策")
    table(d, ["待确认事项", "建议默认值", "原因"], [
        ("首批规则范围", "1-2 个指定期刊或 1 个学校模板 + 通用规则。", "模板差异很大，先保证少量规则真正可验证。"),
        ("论文正文允许的自动修改", "默认仅格式化/转义/交叉引用；所有措辞修改关闭。", "满足“最后版本不能胡编乱造”。"),
        ("题注与图注输入形式", "独立 annotations.zip（Excel/CSV/JSON）+ 可视化确认表单。", "与 [FigN]/[TableN] 对象占位符分离，避免从图片或临近文字推断。"),
        ("链接政策", "界面粘贴 DOI/URL 或 annotations.xlsx 的 Links 表；仅做格式和可访问性检查。", "低门槛批量输入，同时避免自动猜填链接。"),
        ("正式 PDF 定义", "目标官方模板真实编译成功且阻断项为零。", "杜绝降级预览被误用作投稿稿。"),
        ("本地部署边界", "离线优先；联网期刊查询和依赖下载需显式提示。", "论文材料通常包含未发表内容。"),
    ], [4.2, 7.1, 4.9])

    heading(d, "十三、结论与下一步")
    para(d, "项目应以“可信交付”而非“智能补写”为产品核心。最小可行版本的价值在于：让用户把已完成的论文、图表、表格和文献可靠地组装为可编译、可审计的目标格式，而任何缺失内容、歧义匹配或不确定规则都留在可见的人工确认队列。", indent=True)
    para(d, "建议立即进入 P0：确定首批目标模板，收集 5-10 份脱敏样稿及对应资料包，冻结 captions.yaml/links.yaml/manifest.csv 的规范，并将本项目书中的发布门禁转为产品验收用例。", indent=True)

    d.core_properties.title = "论文排版与期刊交付智能体项目计划书"
    d.core_properties.subject = "可行性、需求范围、实施计划与验收标准"
    d.core_properties.author = "Codex"
    d.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
