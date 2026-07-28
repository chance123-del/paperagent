from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "outputs" / "PaperFormat_Agent_使用与验收说明书.docx"

INK = "17211F"
MUTED = "64736D"
MINT = "087A63"
MINT_PALE = "EAF6F1"
LINE = "DCE6E1"
BLUE_PALE = "E8EEF5"


def set_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.margin_top = 80
            cell.margin_bottom = 80


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def style_table(table, header=True):
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if header and row_index == 0:
                set_cell_shading(cell, BLUE_PALE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_font(run, size=9.5, color=INK, bold=(header and row_index == 0))
    if header:
        set_repeat_table_header(table.rows[0])


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, emphasis=None):
    p = doc.add_paragraph(style="Normal")
    if emphasis and emphasis in text:
        before, marked, after = text.partition(emphasis)
        p.add_run(before)
        r = p.add_run(marked)
        set_font(r, size=10.5, color=MINT, bold=True)
        p.add_run(after)
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, MINT_PALE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title + "  ")
    set_font(r, size=10.5, color=MINT, bold=True)
    r = p.add_run(body)
    set_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def fill_table(table, rows):
    for index, values in enumerate(rows):
        row = table.rows[0] if index == 0 else table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = value


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, before, after, color in [
        (1, 16, 18, 10, MINT),
        (2, 13, 14, 7, MINT),
        (3, 12, 10, 5, "1F4D78"),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("PaperFormat Agent  |  使用与验收说明书"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("本地部署版 | 2026-07-28"), size=9, color=MUTED)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("PaperFormat Agent"), size=26, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run("论文格式转换、排版与交付 - 新手使用及验收说明书"), size=14, color=MUTED)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(18)
    set_font(meta.add_run("适用对象：首次使用本地论文排版工具的学生、助教与课题组成员"), size=10, color=MUTED)

    add_note(doc, "交付结论", "已使用一份新建的 Markdown 测试论文完成端到端验收：生成排版工程、规则修复、BibTeX 引用映射、PDF 编译、Word 正式导出均成功。本文档同时记录已发现的使用限制与处理建议。")

    add_heading(doc, "1. 工具用途与适用范围")
    add_body(doc, "PaperFormat Agent 是一个本地运行的论文格式处理工具。它将 Word、PDF、Markdown、LaTeX 或 LaTeX ZIP 项目转换为可编辑的 LaTeX 项目，并生成格式检查报告、PDF 预览及正式交付文件。它适合把已有初稿整理为符合通用论文或期刊/学校规则的排版工程。")
    add_body(doc, "它不是内容代写工具，也不能替代学校或期刊的最终格式审查。特别是封面、授权页、学院专用模板、特殊公式宏包与参考文献细则，必须在最终提交前以目标单位要求为准。", "不能替代学校或期刊的最终格式审查")

    add_heading(doc, "2. 使用前准备")
    add_number(doc, "准备论文原稿。推荐优先使用 .docx、.md 或已有 .tex；若使用 ZIP，请确认压缩包内包含主 .tex 文件和所需图片。")
    add_number(doc, "准备目标格式依据。已有规则可直接选择；没有规则时，可填写学校/期刊名称，上传 Author Guidelines、模板说明或一篇公开参考论文。")
    add_number(doc, "如需自动处理文内数字引用，准备 BibTeX 文献库（.bib）。文内引用请按 [1]、[2,3] 或 [4-6] 的形式写。")
    add_number(doc, "确认本机可访问页面 http://127.0.0.1:7861/。如页面无法打开，双击 web 文件夹中的 run_local_server.vbs 后刷新浏览器。")
    add_note(doc, "文件安全", "原稿、生成项目和编译文件均保存在本机 outputs 目录。不要将含未公开数据或保密附件的文件上传到不受控的第三方服务。")

    add_heading(doc, "3. 页面工作流：从原稿到交付")
    add_heading(doc, "3.1 第一步：上传论文", 2)
    add_body(doc, "在“论文输入”区域二选一操作：点击上传框选择文件，或在“本地论文路径”中粘贴完整路径。不要同时依赖两个来源；优先级以填写的本地路径为准。支持 .docx、.pdf、.md、.markdown、.tex 和 .zip。")
    add_bullet(doc, "Word：适合已有正文与图片的初稿；复杂分栏、文本框、修订痕迹可能需要导出后人工复查。")
    add_bullet(doc, "PDF：仅能可靠提取可选择文本；扫描件需要 OCR，图片文字与复杂表格需额外核对。")
    add_bullet(doc, "Markdown：结构最清晰、最便于测试；表格使用标准竖线表格语法。")
    add_bullet(doc, "LaTeX/ZIP：适合已有项目的格式检查和修复；请保证图片路径与主文件引用一致。")

    add_heading(doc, "3.2 第二步：选择格式要求", 2)
    add_body(doc, "先选择“基础格式规则”。若目标期刊明确在可选期刊规则中，选择对应规则包；否则保持通用规则并上传正式指南。可在“直接填写排版要求”中写出可执行的要求，例如“A4；上下左右边距 2.5 cm；1.5 倍行距；参考文献使用 IEEE”。")
    add_body(doc, "“匹配期刊”用于辅助识别，不应被视为最终依据。若学院提供了 Word/PDF 模板，应上传该文件，并在导出后逐项核对页边距、标题层级、图表题注和参考文献。")

    add_heading(doc, "3.3 第三步：生成快速预览", 2)
    add_body(doc, "选择输出目录后，点击“生成排版工程”。系统会创建一个独立运行目录，生成 source.tex、latex_source.zip、format_report.md 和 compile.log；PDF 预览会在可编译时显示在页面中。快速预览阶段不会直接覆盖你的原稿。")
    add_body(doc, "先阅读格式检查报告，再打开 PDF 预览。重点检查：标题层级、摘要、表格宽度、图题/表题、公式、引用编号、参考文献和页边距。格式评分用于辅助排序，不是可以直接替代人工审核的“投稿通过率”。")

    add_heading(doc, "3.4 第四步：审阅、修订与正式导出", 2)
    add_body(doc, "在 PDF 审阅区域可选中内容作为锚点，用于后续插入公式、图、表或超链接。需要按导师意见修改时，在“在线反馈与二次修订”中输入清晰、可执行的文字，例如“将 Table 1 标题改为中文并置于表格上方”。")
    add_body(doc, "确认预览无误后，点击“正式导出 PDF / Word / LaTeX 源码”。最终应同时保留 PDF、formatted_manuscript.docx、source.tex、formal_latex_source.zip、format_report.md 与 formal_compile.log，便于提交与追溯。")

    add_heading(doc, "4. 输出文件说明")
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    fill_table(table, [
        ("文件", "用途与新手检查点"),
        ("source.tex", "主 LaTeX 源文件。需要模板级调整、摘要环境或特殊命令时，优先修改此文件。"),
        ("latex_source.zip / formal_latex_source.zip", "完整可编辑项目包。提交 LaTeX 源码或交给他人继续修改时使用。"),
        ("main.pdf", "编译得到的 PDF。提交前逐页检查分页、图表、引用和页码。"),
        ("formatted_manuscript.docx", "正式导出的 Word 版本。适合需要 Word 交付的场景，但复杂 LaTeX 排版仍以 PDF/源项目为准。"),
        ("format_report.md", "格式问题、已执行修复和风险评分。先处理 high/critical 项，再看 moderate 项。"),
        ("compile.log / formal_compile.log", "编译记录。出现 PDF 缺失、引用问号或图片找不到时，首先查看此文件。"),
        ("citation_mapping.md", "数字引用与 BibTeX 条目的映射结果。确认每个 [n] 对应预期文献。"),
    ])
    style_table(table)

    add_heading(doc, "5. 新手验收清单")
    add_body(doc, "每次正式提交前，建议按以下顺序复核。任何一项不通过，都不要仅因“PDF 能打开”就直接交付。")
    checklist = [
        "原稿标题、作者、摘要、关键词是否齐全，并与学校/期刊模板要求一致。",
        "PDF 是否能打开，且没有空白页、缺图、表格越界、乱码、问号引用或明显重叠。",
        "格式报告中的 high/critical 项是否已解释或解决；不能解决时是否记录了人工处理决定。",
        "图、表、公式是否都在正文中有提及；图题与表题的位置、语言和编号是否符合目标模板。",
        "文内 [n] 与参考文献是否一一对应；上传 .bib 后是否检查 citation_mapping.md。",
        "页边距、行距、字体、标题层级、参考文献样式是否逐项与目标单位的最新规定核对。",
        "正式交付目录是否同时保留 PDF、Word（如需要）、LaTeX 源码 ZIP、报告和编译日志。",
    ]
    for item in checklist:
        add_bullet(doc, item)

    add_heading(doc, "6. 本次验收实例与结果")
    add_body(doc, "为验证交付链路，本次新建了一篇虚构的 Markdown 论文，内容包含标题、摘要文本、关键词、二级标题、Markdown 表格、LaTeX 敏感字符（%、&、下划线）和 [1] 数字引用，并上传一条 BibTeX 文献记录。测试内容不含真实或未公开研究数据。")
    result_table = doc.add_table(rows=1, cols=3)
    set_table_geometry(result_table, [2300, 3660, 3400])
    fill_table(result_table, [
        ("验收项", "实际结果", "结论"),
        ("输入解析", "Markdown 标题、正文和表格成功转换", "通过"),
        ("规则修复", "应用 thesis_basic；自动修复 3 项", "通过"),
        ("引用映射", "[1] 映射到 smith2024agent；未匹配 0 条", "通过"),
        ("特殊字符", "40%、A&B、metric_score 已转义为 LaTeX 安全写法", "通过"),
        ("PDF 编译", "生成 2 页 main.pdf", "通过"),
        ("正式导出", "PDF 与 Word 均成功；Word 使用内置结构化转换器", "通过"),
        ("格式评分", "77/100 提升至 82/100", "需结合人工审阅"),
    ])
    style_table(result_table)
    add_note(doc, "验收边界", "本次验证的是通用论文规则和本机编译环境。未针对某所学校、某个会议或某本期刊的专用模板做合规承诺；需要提交到具体目标时，仍须上传并核对该目标的官方模板。")

    add_heading(doc, "7. 已发现的小问题与处理方法")
    issues_table = doc.add_table(rows=1, cols=3)
    set_table_geometry(issues_table, [2650, 3610, 3100])
    fill_table(issues_table, [
        ("现象", "原因与影响", "处理建议"),
        ("Markdown 的 ## Abstract 被报告为缺少摘要环境", "当前 Markdown 转换会将 Abstract 变为普通小节，而规则检查要求 LaTeX 的 abstract 环境。PDF 可编译，但严格模板验收可能不通过。", "正式提交前，在 source.tex 中把该小节改为 \\begin{abstract}...\\end{abstract}，或从带摘要环境的 LaTeX 模板开始。"),
        ("首次编译日志出现 undefined citation 警告", "BibTeX 与交叉引用通常需要多轮编译；最终 PDF 已生成，但日志中可能保留前一轮警告。", "以最后一次 formal_compile.log 和 PDF 实际引用显示为准；仍有问号时检查 .bib 键名和正文引用。"),
        ("Word 导出提示未找到 Pandoc", "当前环境使用内置结构化 DOCX 转换器，复杂 LaTeX 宏、公式或特殊表格的 Word 还原度可能低于 PDF。", "以 PDF 作为排版基准；必须交付 Word 时逐页复核，并保留 LaTeX 源项目。"),
        ("格式评分不是满分", "评分反映规则检测到的结构风险，不等同于人工格式审核结果。", "优先处理报告中的 high/critical 项；再按目标模板人工核对封面、摘要、参考文献和图表。"),
        ("扫描 PDF 或复杂 Word 表格效果不稳定", "扫描件依赖 OCR；复杂浮动对象、文本框和嵌套表格的结构信息可能缺失。", "优先提供原始 Word/LaTeX；对 PDF 转换结果逐页比对，必要时手工补录表格与图片。"),
    ])
    style_table(issues_table)

    add_heading(doc, "8. 常见故障排查")
    add_heading(doc, "8.1 页面无法打开", 2)
    add_body(doc, "确认浏览器地址为 http://127.0.0.1:7861/。若出现“拒绝连接”，双击 web/run_local_server.vbs，等待约 5 秒后刷新。仍无法打开时，确认没有安全软件拦截本地 Python 服务，并检查 7861 端口是否被占用。")
    add_heading(doc, "8.2 没有生成 PDF", 2)
    add_body(doc, "打开 compile.log 或 formal_compile.log，搜索 Error、Undefined control sequence、File not found、Citation 等关键词。图片文件丢失时，检查 ZIP 内文件路径；参考文献异常时，检查 .bib 是否为有效 BibTeX 格式及键名是否一致。")
    add_heading(doc, "8.3 PDF 能打开但版式不对", 2)
    add_body(doc, "先确认选择的规则包是否与目标单位一致，再上传官方模板或指南。不要只依赖“匹配期刊”的结果。对于院校封面、中文摘要、目录、致谢、附录、学位授权页等专用内容，应以学校模板为准并人工复核。")
    add_heading(doc, "8.4 怎么把导师意见写得更容易被正确处理", 2)
    add_body(doc, "用“对象 + 动作 + 目标位置/标准”的方式描述。例如：“将表 2 标题改为中文，置于表格上方”；“删除 Results 第 2 段中重复的冒号”；“在 Methods 末尾后插入图 1，图宽 0.72\\linewidth”。避免只写“格式不好看”或“请调整一下”。")

    add_heading(doc, "9. 推荐交付包")
    add_body(doc, "建议为每次论文提交建立一个独立文件夹，包含以下内容：")
    for item in [
        "最终 PDF（提交版本）；",
        "formatted_manuscript.docx（学校或导师要求 Word 时提供）；",
        "formal_latex_source.zip（可复现、可继续编辑的源码项目）；",
        "format_report.md 与 formal_compile.log（审计与故障追溯）；",
        "目标单位的原始模板/指南副本，以及一份人工核对记录。",
    ]:
        add_bullet(doc, item)
    add_note(doc, "最后提醒", "在校内提交、送审或投稿前，请把官方模板逐项作为最终依据。工具负责提高转换效率与可追溯性，最终格式责任仍应由提交人完成核对。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "PaperFormat Agent 使用与验收说明书"
    doc.core_properties.subject = "本地论文格式转换工具的新手操作与验收指南"
    doc.core_properties.author = "PaperFormat Agent"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
